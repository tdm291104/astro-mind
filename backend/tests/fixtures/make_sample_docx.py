"""Generate tests/fixtures/sample.docx with known headings and content.

Run once from the repo root: `uv run python tests/fixtures/make_sample_docx.py`.
The output is committed to git so tests do not regenerate it. Facts mirror
sample.pdf so the two fixtures can be cross-checked in multi-doc tests.
"""
from pathlib import Path

from docx import Document

HERE = Path(__file__).parent
OUT = HERE / "sample.docx"

# (heading or None for preamble, [body paragraphs])
SECTIONS = [
    (None, [
        "This is a synthetic DOCX fixture for the AstroMind test suite.",
        "It is not derived from any external source.",
    ]),
    ("The Sun", [
        "The Sun is the star at the center of the Solar System.",
        "Its core temperature is approximately 15 million kelvin.",
    ]),
    ("Jupiter", [
        "Jupiter is the largest planet in the Solar System.",
        "Its mass is approximately 318 Earth masses.",
    ]),
    ("Exoplanets", [
        "An exoplanet is a planet outside our Solar System.",
        "TRAPPIST-1e is a confirmed exoplanet, announced in February 2017.",
    ]),
]


def main() -> None:
    doc = Document()
    for heading, body in SECTIONS:
        if heading is not None:
            doc.add_heading(heading, level=1)
        for para in body:
            doc.add_paragraph(para)
    doc.save(str(OUT))
    n_headings = sum(1 for h, _ in SECTIONS if h is not None)
    print(f"Wrote {OUT} ({n_headings} headings + preamble)")


if __name__ == "__main__":
    main()
